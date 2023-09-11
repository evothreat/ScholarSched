from tkinter import *
from tkinter.filedialog import askopenfilename
from tkinter.scrolledtext import ScrolledText
import csv
import os
from os.path import exists, basename, join as pathjoin, dirname
from shutil import rmtree
from docx import Document


WORKSPACE = 'wsp'
BSO_DIR = pathjoin(WORKSPACE, 'BSO')
EDITED_DIR = pathjoin(WORKSPACE, 'Edited')
CERTS_DIR = pathjoin(WORKSPACE, 'Zertifikate')

def join_dicts(*ds):
    d = {}
    for d2 in ds:
        d.update(d2)
    return d

class Backend:
    def __init__(self):
        if not exists(WORKSPACE):
            os.mkdir(WORKSPACE)

        if not exists(EDITED_DIR):
            os.mkdir(EDITED_DIR)

        if not exists(CERTS_DIR):
            os.mkdir(CERTS_DIR)

        self.root = Tk()
        self.window = Window(self.root, self)
        self.root.update()
        self.root.mainloop()

    def create_filepaths(self):
        """
        ->  erstellt ein Verzeichnis und eine Datei für jeden Schüler
            aus der gewählten Datei

        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        c = Counter()
        with open(rfile) as rf:
            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                if not exists(m.fpath):
                    os.makedirs(dirname(m.fpath), exist_ok=True)

                    with open(m.fpath, 'w') as wf:
                        w = csv.writer(wf, delimiter=';')
                        w.writerow([m.lname, m.fname, m.teacher])
                        w.writerow(Heading.event())
                        c.inc_succes()                        
                    self.window.write('[+] Dateipfad für', m.lname, m.fname,
                                      'erstellt', tag='succes')
                    continue

                with open(m.fpath) as rf:
                    row2 = next(csv.reader(rf, delimiter=';'))
                    m2 = Member(
                        lname=row2[0], 
                        fname=row2[1], 
                        teacher=row2[2]
                    )
                    if m != m2:
                        c.inc_error()
                        self.window.write('[-] Unterschiedliche Daten bei',
                                          m.lname, m.fname, tag='error')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def delete_filepaths(self):
        """
        ->  löscht das Verzeichnis und die Datei des jeweiligen 
            Schülers aus der gewählten CSV-Datei

        """
        self.window.flush_output()
        
        rfile = askopenfilename()
        if not rfile:
            return

        c = Counter()
        with open(rfile) as rf:
            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                if not exists(m.id):
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')
                    continue

                rmtree(m.id)
                c.inc_succes()
                self.window.write('[+] Dateipfad von', m.lname, m.fname,
                                  'entfernt', tag='succes')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def assign_id(self):
        """
        ->  ordnet jedem Schüler aus der gewählten CSV-Datei
            eine Schülernummer zu und speichert dies in einer
            CSV-Datei im Verzeichnis 'Edited'

        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        mbrs = []
        for id in os.listdir(BSO_DIR):
            m = Member(id=id)
            with open(m.fpath) as rf:
                row = next(csv.reader(rf, delimiter=';'))
                m.lname = row[0]
                m.fname = row[1]
                m.teacher = row[2]
                mbrs.append(m)

        wfile = pathjoin(EDITED_DIR, basename(rfile))

        c = Counter()
        with open(rfile) as rf, open(wfile, 'w', newline='') as wf:
            w = csv.DictWriter(wf, Heading.template(), delimiter=';')
            w.writeheader()
            w.writerow({k: '' for k in Heading.template()})

            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m.fname or not m.lname:
                    continue

                # or iterate simply through mbrs
                for i in range(len(mbrs)):
                    if m == mbrs[i]:
                        m2 = mbrs.pop(i)
                        m.id = m2.id
                        m.teacher = m2.teacher
                        c.inc_succes()
                        self.window.write('[+] SN und Tutor von', m.lname, m.fname, 
                                          'gefunden', tag='succes')
                        break
                else:
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')

                e = event(row)
                w.writerow(join_dicts(e.dict(), m.dict()))

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def update_teacher(self):
        """
        ->  aktualisiert den Tutor in der Datei des jeweiligen Schülers
            aus der gewählten CSV-Datei

        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        c = Counter()
        with open(rfile) as rf:
            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                if not exists(m.fpath):
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')
                    continue

                rows = []
                with open(m.fpath) as rf2:
                    rows = list(csv.reader(rf2, delimiter=';'))

                # index of old teacher
                if m.teacher == rows[0][2]:
                    continue
                rows[0][2] = m.teacher

                with open(m.fpath, 'w', newline='') as wf:
                    csv.writer(wf, delimiter=';').writerows(rows)

                c.inc_succes()
                self.window.write('[+] Tutor von', m.lname, m.fname,
                                  'aktualisiert', tag='succes')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def add_events(self):
        """
        ->  trägt die Veranstaltungen aus der gewählten CSV-Datei
            in die jeweilige Schülerdatei ein

        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        c = Counter()
        with open(rfile) as rf:
            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                if not exists(m.fpath):
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')
                    continue

                e = event(row)
                with open(m.fpath, 'a', newline='') as wf:
                    w = csv.writer(wf, delimiter=';')
                    w.writerow(e.dict().values())

                c.inc_succes()
                self.window.write('[+] Eintrag für', m.lname, m.fname,
                                  'gemacht', tag='succes')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def add_single_event(self):
        """
        ->  trägt die Veranstaltung aus der Eingabe in
            die entprechende Datei ein

        """
        self.window.flush_output()

        m = Member(id=self.window.input[0].get())
        e = Event(
            name=self.window.input[1].get(),
            date=self.window.input[2].get(),
            module=self.window.input[3].get(),
            points=self.window.input[4].get()
        )

        if not e:
            self.window.write('[-] Überprüfen Sie die Eingaben', tag='error')
            return 

        if not exists(m.fpath):
            self.window.write('[-] SchülerIn existiert nicht', tag='error')
            return

        with open(m.fpath) as rf:
            row = next(csv.reader(rf, delimiter=';'))
            m.lname = row[0]
            m.fname = row[1]
            m.teacher = row[2]

        with open(m.fpath, 'a', newline='') as wf:
            w = csv.writer(wf, delimiter=';')
            w.writerow(e.dict().values())

        self.window.write('[+] Eintrag für', m.lname, m.fname,
                          'gemacht', tag='succes')

    def get_events(self):
        """
        ->  zeigt die Veranstaltungen des gewünschten Schülers

        """
        self.window.flush_output()

        m = Member(id=self.window.input[0].get())

        if not exists(m.fpath):
            self.window.write('[-] SchülerIn existiert nicht', tag='error')
            return

        with open(m.fpath) as rf:
            r = csv.reader(rf, delimiter=';')

            m.lname, m.fname, m.teacher = next(r)

            next(r)
            next(r)
            next(r)

            self.window.write('[+] Eintraege für', m.lname, m.fname,
                              tag='succes')

            for row in r:
                e = Event(
                        name=row[0],
                        date=row[1],
                        module=row[2],
                        points=row[3]
                )
                self.window.write('=> Veranstaltung:', e.name,  '\n'
                                  '=> Datum:        ', e.date,  '\n'
                                  '=> Modul:        ', e.module,'\n'
                                  '=> BSO-Punkte:   ', e.points)

    def calculate_points(self):
        """
        ->  berechnet die BSO-Punkte für den entsprechenden Schüler
            aus der gewählten CSV-Datei und speichert dies
            in einer CSV-Datei im Verzeichnis 'Edited'

        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        wfile = pathjoin(EDITED_DIR, basename(rfile))

        c = Counter()
        with open(rfile) as rf, open(wfile, 'w', newline='') as wf:
            w = csv.DictWriter(wf, Heading.calc(), delimiter=';')
            w.writeheader()
            w.writerow({k: '' for k in Heading.calc()})

            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                if not exists(m.fpath):
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')
                    continue

                with open(m.fpath) as rf2:
                    r = csv.reader(rf2, delimiter=';')
                    next(r)
                    next(r)
                    next(r)
                    next(r)
                    for row2 in r:
                        e = Event(
                                name=row2[0],
                                date=row2[1],
                                module=row2[2],
                                points=row2[3]
                        )

                        try:
                            m.info[e.module] += int(e.points)
                        except (KeyError, ValueError):
                            continue

                    m.calculate()
                    w.writerow(join_dicts(m.dict(), m.info))
                    
                    c.inc_succes()
                    self.window.write('[+] Punkte für', m.lname, m.fname,
                                      'berechnet', tag='succes')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())

    def create_certs(self):
        """
        ->  überprüft ob der jeweilige Schüler aus der gewählten
            CSV-Datei ein Zertifikat verdient hat und falls ja
            ruft die Funktion 'cert()' auf, die das Zertifikat
            für den Schüler erstellt und im Verzeichnis
            'Zertifikate' absichert
        """
        self.window.flush_output()

        rfile = askopenfilename()
        if not rfile:
            return

        c = Counter()
        with open(rfile) as rf:
            for row in csv.DictReader(rf, delimiter=';'):
                m = member(row)
                if not m:
                    continue

                # or try except or if
                for k in m.info:
                    m.info[k] = row.get(k, 0)

                if not m.certified():
                    continue

                if not exists(m.fpath):
                    c.inc_error()
                    self.window.write('[-] SchülerIn', m.lname, m.fname,
                                      'existiert nicht', tag='error')
                    continue

                with open(m.fpath) as rf2:
                    r = csv.reader(rf2, delimiter=';')
                    next(r)
                    next(r)
                    next(r)
                    next(r)

                    evts = []                  
                    for row2 in r:
                        evts.append({
                            'EVENT':  row2[0],
                            'DATE':   row2[1],
                            #'MODULE': row[2],
                            'POINTS': row2[3]
                        })

                    templ = Template('template.docx')

                    templ.replace_str('NAME',   m.fullname())
                    templ.replace_str('ID',     m.id)
                    templ.replace_str('TOTAL',  m.total_pts())      # change?
                    templ.update_table(evts)

                    templ.save(pathjoin(CERTS_DIR, m.id + '.docx'))

                    c.inc_succes()
                    self.window.write('[+] Zertifikat für', m.lname, m.fname,
                                      'erstellt', tag='succes')

        self.window.write('[*] Erfolge:', c.get_succes(), '\n'
                          '[*] Fehler: ', c.get_error())
                            
class Window:
    def __init__(self, master, backend):
        self.master = master

        self.backend = backend

        self.master.title('ScholarSched')
        self.master.geometry('820x630')

        self.bar = Menu()
        self.master.config(menu=self.bar)

        self.lists = Menu(self.bar, tearoff=False)
        self.help = Menu(self.bar, tearoff=False)

        self.lists.add_command(label='Schülerverzeichnisse anlegen',
                               command=self.backend.create_filepaths)
        self.lists.add_command(label='Schülernummer zuordnen',
                               command=self.backend.assign_id)
        self.lists.add_command(label='Veranstaltungen eintragen',
                               command=self.backend.add_events)
        self.lists.add_command(label='BSO-Punkte berechnen',
                               command=self.backend.calculate_points)
        self.lists.add_command(label='Zertifikate erstellen',
                               command=self.backend.create_certs)
        self.lists.add_command(label='Tutor aktualisieren',
                                command=self.backend.update_teacher)
        self.lists.add_command(label='Schülerverzeichnisse entfernen',
                               command=self.backend.delete_filepaths)

        self.help.add_command(label='Schülerverzeichnisse anlegen',
                              command=lambda: os.system('start doku/p01.pdf'))
        self.help.add_command(label='Schülernummer zuordnen',
                              command=lambda: os.system('start doku/p02.pdf'))
        self.help.add_command(label='Veranstaltungen eintragen',
                              command=lambda: os.system('start doku/p03.pdf'))
        self.help.add_command(label='BSO-Punkte berechnen',
                              command=lambda: os.system('start doku/p04.pdf'))
        self.help.add_command(label='Zertifikate erstellen',
                              command=lambda: os.system('start doku/p05.pdf'))
        self.help.add_command(label='Tutor aktualisieren',
                              command=lambda: os.system('start doku/p06.pdf'))
        self.help.add_command(label='Schülerverzeichnisse entfernen',
                              command=lambda: os.system('start doku/p07.pdf'))

        self.bar.add_cascade(label='Listenoperationen', menu=self.lists)
        self.bar.add_cascade(label='Hilfe', menu=self.help)

        self.headings = [Label(self.master, text='Schülernummer'),
                         Label(self.master, text='Veranstaltung'),
                         Label(self.master, text='Datum'),
                         Label(self.master, text='Modul'),
                         Label(self.master, text='BSO-Punkte')]

        self.input = [StringVar() for i in range(5)]
        self.entries = [Entry(self.master, textvariable=v) for v in self.input]

        sticky = W + E + N + S

        for i in range(5):
            self.headings[i].grid(row=0, column=i, sticky=sticky)
            self.entries[i].grid(row=1, column=i, sticky=sticky)

        self.search = Button(self.master, text='Suchen',
                             command=self.backend.get_events)
        self.add = Button(self.master, text='Eintragen',
                          command=self.backend.add_single_event)

        self.search.grid(row=2, column=0, columnspan=1, sticky=sticky)
        self.add.grid(row=2, column=1, columnspan=4, sticky=sticky)

        self.logo = PhotoImage(file='logo/cvo-gyo-logo.png')

        self.lab = Label(self.master, image=self.logo)
        self.lab.grid(row=0, column=5, rowspan=3, sticky=sticky)

        self.output = ScrolledText(self.master, width=100, height=35)
        self.output.grid(row=3, columnspan=6, sticky=sticky)

        self.output.tag_add('succes', '1.0', END)
        self.output.tag_add('error', '1.0', END)

        self.output.tag_config('succes', foreground='green')
        self.output.tag_config('error', foreground='red')

    def write(self, *msg, tag=''):
        self.output.config(state=NORMAL)
        self.output.insert(END, ' '.join(msg) + '\n\n', tag)
        self.output.config(state=DISABLED)

    def flush_output(self):
        self.output.config(state=NORMAL)
        self.output.delete(1.0, END)
        self.output.config(state=DISABLED)


class Member:
    def __init__(self, id='', fname='', lname='', teacher=''):
        self.id = id
        self.fname = fname
        self.lname = lname
        self.teacher = teacher
        self.info = {
            'M1':         0,
            'M2':         0,
            'M3':         0,
            'M4':         0,
            'M5':         0,
            'Gesamt':     0,
            'Zertifikat': 'nein'
        }
        self.fpath = pathjoin(BSO_DIR, id, id + '.csv')

    def fullname(self):
        return self.fname + ' ' + self.lname

    def total_pts(self):
        return self.info['Gesamt']

    def calculate(self):
        self.info['Gesamt'] += self.info['M1'] + self.info['M2'] + \
                               self.info['M3'] + self.info['M4'] + \
                               self.info['M5']

        if self.info['M1'] >= 20 and \
           self.info['M2'] >= 10 and \
           self.info['M3'] >= 5  and \
           self.info['M4'] >= 15 and \
           self.info['M5'] >= 10:
            self.info['Zertifikat'] = 'ja'

    def certified(self):
        return self.info['Zertifikat'] == 'ja'

    def dict(self):
        return {
            'Vorname':  self.fname,
            'Nachname': self.lname,
            'SN':       self.id,
            'Tutor':    self.teacher
        }

    def __bool__(self):
        return bool(self.id and self.fname and self.lname)

    def __eq__(self, other):
        return self.fname == other.fname and \
               self.lname == other.lname

    def __ne__(self, other):
        return self.fname != other.fname or  \
               self.lname != other.lname

class Event:
    def __init__(self, name='', date='', module='', points=''):
        self.name = name
        self.date = date
        self.module = module
        self.points = points

    def dict(self):
        return {
            'Veranstaltung': self.name,
            'Datum':         self.date,
            'Modul':         self.module,
            'BSO-Punkte':    self.points
        }

    def __bool__(self):
        return bool(self.name and self.date and \
                    self.module and self.points)

class Template:
    def __init__(self, name):
        self.doc = Document(name)

    def replace_str(self, old, new, obj=None):
        if not obj:
            obj = self.doc
        for p in obj.paragraphs:
            if old in p.text:           # no need but faster
                for r in p.runs:
                    if old in r.text:
                        r.text = r.text.replace(old, new)
                        return

    def update_table(self, values):
        keys = values[0].keys()
        for t in self.doc.tables:
            for row in t.rows:
                new = values.pop() if values else {}
                for c in row.cells:
                    for k in keys:
                        if k in c.text:
                            self.replace_str(k, new.get(k, ''), obj=c)
                            break
                    else:
                        values.append(new)
                        break

    def save(self, name):
        self.doc.save(name)

class Counter:
    def __init__(self):
        self.succes = 0
        self.error = 0

    def inc_succes(self):
        self.succes += 1

    def inc_error(self):
        self.error += 1

    def get_succes(self):
        return str(self.succes)

    def get_error(self):
        return str(self.error)

class Heading:
    def event():
        return (
            'Veranstaltung',
            'Datum', 'Modul',
            'BSO-Punkte'
        )

    def template():
        return (
            'Veranstaltung', 'Datum',
            'Nachname', 'Vorname',
            'SN', 'Tutor',
            'Modul', 'BSO-Punkte'
        )

    def calc():
        return (
            'Nachname', 'Vorname',
            'SN', 'Tutor', 
            'M1', 'M2', 'M3', 'M4', 'M5',
            'Gesamt', 'Zertifikat'
        )

def member(d):
    return Member(
        d.get('SN', ''),
        d.get('Vorname', ''),
        d.get('Nachname', ''),
        d.get('Tutor', '')
    )

def event(d):
    return Event(
        d.get('Veranstaltung', ''),
        d.get('Datum', ''),
        d.get('Modul', ''),
        d.get('BSO-Punkte', '')
    )

b = Backend()

# print counter status
